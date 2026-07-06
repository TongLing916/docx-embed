from __future__ import annotations

import base64
import json
from pathlib import Path

from docx import Document

from edp.main_parser import parse_main_document
from examples.run_parser import run_parser


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


def test_ragflow_parser_uses_markdown_export(tmp_path: Path) -> None:
    source = _write_ragflow_docx(tmp_path / "ragflow.docx")

    result = parse_main_document(source, "ragflow", tmp_path / "artifacts")

    assert result.parser == "ragflow"
    assert result.warnings == []
    assert result.artifacts["clean_markdown"] == tmp_path / "artifacts" / "ragflow.md"
    assert "Revenue Section\n===============" in result.markdown
    assert "Intro paragraph" in result.markdown
    assert "<table>" not in result.markdown
    assert "data:image" not in result.markdown
    assert "base64" not in result.markdown
    assert "![image_001](images/image_001.png)" in result.markdown
    assert (tmp_path / "artifacts" / "images" / "image_001.png").exists()


def test_pure_ragflow_mode_writes_package_and_materializes_images(tmp_path: Path) -> None:
    source = _write_ragflow_docx(tmp_path / "pure-ragflow.docx")

    exit_code = run_parser(source, tmp_path / "package", "ragflow")

    assert exit_code == 0
    content = (tmp_path / "package" / "structured" / "content.md").read_text(encoding="utf-8")
    assert "Intro paragraph" in content
    assert "![image_001](assets/images/ragflow/images/image_001.png)" in content
    assert (
        tmp_path / "package" / "structured" / "assets" / "images" / "ragflow" / "images" / "image_001.png"
    ).exists()

    manifest = json.loads((tmp_path / "package" / "manifest.json").read_text(encoding="utf-8"))
    assert manifest["content_map"]["parser"] == "ragflow"


def test_ragflow_markdown_parser_name_is_unsupported(tmp_path: Path) -> None:
    source = _write_ragflow_docx(tmp_path / "ragflow-markdown.docx")

    result = parse_main_document(source, "ragflow-markdown", tmp_path / "markdown-artifacts")

    assert result.parser == "ragflow-markdown"
    assert result.markdown == ""
    assert result.warnings == ["Unsupported main parser: ragflow-markdown"]


def test_ragflow_parser_materializes_table_cell_images_from_markdown_export(tmp_path: Path) -> None:
    source = _write_ragflow_docx_with_table_image(tmp_path / "table-image.docx")

    result = parse_main_document(source, "ragflow", tmp_path / "artifacts")

    assert result.warnings == []
    assert "| Logo image_001 |" in result.markdown
    assert (tmp_path / "artifacts" / "images" / "image_001.png").exists()


def test_pure_ragflow_mode_materializes_table_cell_images(tmp_path: Path) -> None:
    source = _write_ragflow_docx_with_table_image(tmp_path / "pure-table-image.docx")

    exit_code = run_parser(source, tmp_path / "package", "ragflow")

    assert exit_code == 0
    content = (tmp_path / "package" / "structured" / "content.md").read_text(encoding="utf-8")
    assert "| Logo image_001 |" in content


def _write_ragflow_docx(path: Path) -> Path:
    document = Document()
    document.add_heading("Revenue Section", level=1)
    document.add_paragraph("Intro paragraph")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "1200"
    image_path = path.with_name("image.png")
    image_path.write_bytes(PNG_1X1)
    document.add_picture(str(image_path))
    caption = document.add_paragraph("Figure caption")
    caption.style = "Caption"
    document.add_paragraph("Outro paragraph")
    document.save(path)
    return path


def _write_ragflow_docx_with_table_image(path: Path) -> Path:
    document = Document()
    document.add_heading("Asset Section", level=1)
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "Logo "
    image_path = path.with_name("table-image.png")
    image_path.write_bytes(PNG_1X1)
    cell.paragraphs[0].add_run().add_picture(str(image_path))
    document.save(path)
    return path
